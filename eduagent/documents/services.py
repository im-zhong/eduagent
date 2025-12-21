from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter

from eduagent.documents.models import DocumentIngestionJob
from eduagent.documents.repository import DocumentRepository
from eduagent.llm.factory import get_embedding_model
from eduagent.storage.milvus_store import (
    EmbeddingRecord,
    MilvusVectorStore,
    milvus_store,
)


class DocxIngestionService:
    """Parses DOCX files, chunks them, and stores chunks via the repository."""

    def __init__(
        self,
        repository: DocumentRepository,
        *,
        chunk_size: int = 800,
        chunk_overlap: int = 100,
    ) -> None:
        self.repository = repository
        # Custom splitter: LangChain's defaults only break on ASCII punctuation, so
        # long Chinese runs blew past Milvus' 2048-char VARCHAR cap. The rewritten
        # splitter below includes common Chinese delimiters (。, ，, ；, …) plus an
        # empty-string fallback so we always hard-wrap before indexing.
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,  # tokens ≠ chars; Chinese chars are dense
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n",
                "\n",
                "。",
                "！",
                "？",  # Chinese sentence end
                "；",
                "，",
                " ",  # fallback
                "",
            ],
        )

    def _load_text(self, file_path: str | Path) -> list[str]:
        doc = DocxDocument(str(Path(file_path)))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Table-heavy documents often keep their content in table cells instead of
        # top-level paragraphs, so collect that text as well.
        table_text: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        table_text.append(text)
        combined = paragraphs + table_text
        if not combined:
            msg = f"No readable text found in {file_path}"
            raise ValueError(msg)
        return combined

    def _chunk_paragraphs(self, paragraphs: list[str]) -> list[LCDocument]:
        base_doc = LCDocument(
            page_content="\n\n".join(paragraphs),
            metadata={"paragraphs": len(paragraphs)},
        )
        return self.splitter.split_documents([base_doc])

    async def ingest_docx(
        self,
        *,
        source_filename: str,
        file_path: str,
        subject: str | None,
        grade_level: str | None,
        metadata: dict[str, Any] | None = None,
    ) -> DocumentIngestionJob:
        job = await self.repository.create_job(
            source_filename=source_filename,
            file_path=file_path,
            subject=subject,
            grade_level=grade_level,
            metadata=metadata,
        )
        try:
            paragraphs = self._load_text(file_path)
            chunks = self._chunk_paragraphs(paragraphs)
            total_paragraphs = len(paragraphs)
            for index, chunk in enumerate(chunks):
                chunk_metadata = {
                    "paragraphs": total_paragraphs,
                    "chunk_index": index,
                    "character_count": len(chunk.page_content),
                }
                await self.repository.add_chunk(
                    job.id,
                    chunk_index=index,
                    content=chunk.page_content,
                    token_count=len(chunk.page_content.split()),
                    extras={
                        "metadata": chunk_metadata,
                    },
                )
        except Exception as exc:  # pragma: no cover - defensive
            await self.repository.update_status(
                job.id,
                status="failed",
                error_message=str(exc),
            )
            raise
        else:
            updated = await self.repository.update_status(
                job.id,
                status="completed",
                total_chunks=len(chunks),
            )
            return updated or job


class EmbeddingBackend:
    """Lightweight wrapper to make embedding dependencies swappable."""

    BATCH_LIMIT = 64

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        if not texts:
            return []
        model = get_embedding_model()
        results: list[list[float]] = []
        batch_limit = max(1, self.BATCH_LIMIT)
        for start in range(0, len(texts), batch_limit):
            batch = texts[start : start + batch_limit]
            batch_vectors = model.embed_documents(batch)
            results.extend(batch_vectors)
        return results

    def embed_query(self, text: str) -> list[float]:
        model = get_embedding_model()
        return model.embed_query(text)


class ChunkEmbeddingService:
    """Generates embeddings for chunks and indexes them in Milvus."""

    def __init__(
        self,
        repository: DocumentRepository,
        *,
        vector_store: MilvusVectorStore | None = None,
        embedder: EmbeddingBackend | None = None,
    ) -> None:
        self.repository = repository
        self.vector_store = vector_store or milvus_store
        self.embedder = embedder or EmbeddingBackend()

    async def index_job_chunks(self, job_id: str) -> int:
        chunks = await self.repository.list_chunks(job_id)
        if not chunks:
            return 0
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedder.embed_documents(texts)
        print(f"embeder: f{self.embedder}")
        if len(embeddings) != len(chunks):
            msg = "Embedding backend returned mismatched vector count"
            raise ValueError(msg)

        def _clip(text: str) -> str:
            limit = getattr(self.vector_store, "text_limit", 2048)
            return text[:limit]

        records = [
            EmbeddingRecord(
                record_id=chunk.id,
                text=_clip(chunk.content),
                embedding=vector,
                metadata={
                    **chunk.chunk_metadata,
                    "ingestion_job_id": chunk.ingestion_job_id,
                },
            )
            for chunk, vector in zip(chunks, embeddings, strict=True)
        ]
        # for record in records:
        #     assert len(record.embedding) == 2048
        inserted = self.vector_store.insert_records(records)
        for chunk in chunks:
            await self.repository.set_chunk_vector_id(chunk.id, vector_id=chunk.id)
        await self.repository.update_status(job_id, status="indexed")
        return inserted

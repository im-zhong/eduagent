from __future__ import annotations

import asyncio
from collections.abc import Generator, Sequence
from pathlib import Path
from typing import Any, Protocol, cast

import pytest
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from eduagent.api.schemas import (
    QuestionGenerationResponse,
    QuizAnswerItem,
    QuizEvaluationPayload,
    QuizEvaluationResponse,
    QuizGenerationPayload,
    QuizGenerationRules,
    SubjectArea,
    TextbookIngestionResult,
    TextbookUploadMetadata,
)
from eduagent.quiz.enums import JobStatus
from eduagent.quiz.repository import QuizJobRepository
from eduagent.storage.database_service import DatabaseService, QuizJobCreate
from eduagent.storage.engine import async_engine
from eduagent.storage.minio_service import minio_service
from eduagent.tasks import quiz as quiz_tasks
from eduagent.user.models import Base


class SupportsResult(Protocol):
    def get(
        self,
        timeout: float | None = ...,
        *,
        propagate: bool = ...,
        disable_sync_subtasks: bool = ...,
        **kwargs: object,
    ) -> object: ...


class SupportsApply(Protocol):
    def apply(
        self,
        args: Sequence[object] | None = ...,
        kwargs: dict[str, object] | None = ...,
    ) -> SupportsResult: ...


def _apply_task(task: object, *args: object) -> object:
    celery_task = cast(SupportsApply, task)
    applied = celery_task.apply(args=args)
    return applied.get()


@pytest.fixture
def sqlite_sessionmaker(
    monkeypatch: pytest.MonkeyPatch,
) -> Generator[async_sessionmaker[AsyncSession]]:
    engine = create_async_engine("sqlite+aiosqlite://", future=True)

    async def _prepare() -> None:
        async with engine.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)

    asyncio.run(_prepare())
    session_factory = async_sessionmaker(
        engine, class_=AsyncSession, expire_on_commit=False
    )
    monkeypatch.setattr(quiz_tasks, "task_session_factory", session_factory)

    yield session_factory

    asyncio.run(engine.dispose())


def _create_job(
    session_factory: async_sessionmaker[AsyncSession],
    *,
    subject: str,
    grade_level: str,
) -> str:
    async def _create() -> str:
        async with session_factory() as session:
            repo = QuizJobRepository(session)
            job = await repo.create_ingestion_job(
                source_filename="sample.docx",
                file_path="sample.docx",
                subject=subject,
                grade_level=grade_level,
                payload={},
            )
            return job.id

    return asyncio.run(_create())


def _get_job_status(
    session_factory: async_sessionmaker[AsyncSession], job_id: str
) -> tuple[str, dict[str, Any]]:
    async def _load() -> tuple[str, dict[str, object]]:
        async with session_factory() as session:
            repo = QuizJobRepository(session)
            job = await repo.get_job(job_id)
            assert job is not None
            return job.status, job.result_payload

    return asyncio.run(_load())


def test_generate_quiz_updates_job(
    sqlite_sessionmaker: async_sessionmaker[AsyncSession],
) -> None:
    job_id = _create_job(sqlite_sessionmaker, subject="Physics", grade_level="10")
    total_questions = 4
    payload = QuizGenerationPayload(
        job_id=job_id,
        subject=SubjectArea.SCIENCE,
        rules=QuizGenerationRules(total_questions=total_questions),
    )

    result_raw = quiz_tasks.generate_quiz(job_id, payload)
    result = QuestionGenerationResponse.model_validate(result_raw)
    assert len(result.questions) == total_questions
    status, stored_result = _get_job_status(sqlite_sessionmaker, job_id)
    assert status == JobStatus.COMPLETED.value
    assert stored_result["questions"]
    assert stored_result["generation_id"] == job_id


def test_evaluate_answers_updates_job(
    sqlite_sessionmaker: async_sessionmaker[AsyncSession],
) -> None:
    job_id = _create_job(sqlite_sessionmaker, subject="History", grade_level="11")
    payload = QuizEvaluationPayload(
        job_id=job_id,
        answers=[
            QuizAnswerItem(question_id="q1", answer="A", is_correct=True),
            QuizAnswerItem(question_id="q2", answer="B", is_correct=False),
        ],
    )

    result_raw = quiz_tasks.evaluate_answers(job_id, payload)
    result = QuizEvaluationResponse.model_validate(result_raw)
    assert result.score == 1
    status, stored_result = _get_job_status(sqlite_sessionmaker, job_id)
    assert status == JobStatus.COMPLETED.value
    assert stored_result["total"] == len(payload.answers)


def test_generate_quiz_task_apply_returns_response(
    sqlite_sessionmaker: async_sessionmaker[AsyncSession],
) -> None:
    job_id = _create_job(sqlite_sessionmaker, subject="Chemistry", grade_level="11")
    total_questions = 2
    payload = QuizGenerationPayload(
        job_id=job_id,
        subject=SubjectArea.SCIENCE,
        rules=QuizGenerationRules(total_questions=total_questions),
    )
    task_result = _apply_task(quiz_tasks.generate_quiz, job_id, payload)
    response = QuestionGenerationResponse.model_validate(task_result)
    assert len(response.questions) == total_questions


def test_evaluate_answers_task_apply_returns_score(
    sqlite_sessionmaker: async_sessionmaker[AsyncSession],
) -> None:
    job_id = _create_job(sqlite_sessionmaker, subject="Geography", grade_level="9")
    payload = QuizEvaluationPayload(
        job_id=job_id,
        answers=[
            QuizAnswerItem(question_id="q1", answer="A", is_correct=True),
            QuizAnswerItem(question_id="q2", answer="B", is_correct=False),
        ],
    )
    task_result = _apply_task(quiz_tasks.evaluate_answers, job_id, payload)
    response = QuizEvaluationResponse.model_validate(task_result)
    assert response.score == 1


def _create_sample_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Sample material for ingestion tests.")
    doc.add_paragraph("Second paragraph with meaningful text.")
    doc.save(str(path))


def test_process_upload_task_apply_indexes_document(tmp_path: Path) -> None:
    doc_path = tmp_path / "lesson.docx"
    _create_sample_docx(doc_path)
    with doc_path.open("rb") as file_obj:
        stored = minio_service.store_file(
            file_obj,
            filename=doc_path.name,
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
    metadata = TextbookUploadMetadata(
        filename=stored.object_name,
        original_filename=doc_path.name,
        subject=SubjectArea.SCIENCE,
        grade_level="Undergraduate",
    )
    db_service = DatabaseService()
    job_record = db_service.create_quiz_job(
        QuizJobCreate(
            source_filename=metadata.filename,
            file_path=str(doc_path),
            subject=metadata.subject.value if metadata.subject else None,
            grade_level=metadata.grade_level,
            job_payload={"filename": metadata.filename},
        )
    )
    document_job_id: str | None = None
    try:
        asyncio.run(async_engine.dispose())
        task_result = _apply_task(
            quiz_tasks.process_textbook_upload,
            job_record.id,
            stored.object_name,
            metadata,
        )
        ingestion = TextbookIngestionResult.model_validate(task_result)
        assert ingestion.chunks >= 1
        assert ingestion.embedded_records == ingestion.chunks
        document_job_id = ingestion.document_job_id
        assert document_job_id
    finally:
        db_service.delete_quiz_job(job_record.id)
        if document_job_id:
            db_service.delete_document_job(document_job_id)

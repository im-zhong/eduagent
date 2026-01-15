from __future__ import annotations

from pathlib import Path

import pypandoc
import pytest

from docx import Document

from eduagent.parser.parser_service import (
    chunk_markdown,
    convert_document_to_markdown,
    split_markdown,
)


def test_split_markdown_splits_on_blank_lines() -> None:
    markdown = "First paragraph.\n\nSecond paragraph.\n\n\nThird paragraph."
    chunks = split_markdown(markdown)
    assert chunks == [
        "First paragraph.",
        "Second paragraph.",
        "Third paragraph.",
    ]


def test_split_markdown_ignores_empty_input() -> None:
    assert split_markdown("   \n\n") == []


def _assert_pandoc_available() -> None:
    _ = pypandoc.get_pandoc_version()


@pytest.mark.anyio
async def test_convert_document_from_markdown_file(tmp_path: Path) -> None:
    markdown_path = tmp_path / "sample.md"
    markdown_path.write_text("Alpha paragraph.\n\nBeta paragraph.\n")

    markdown = await convert_document_to_markdown(markdown_path)
    chunks = chunk_markdown(markdown)

    assert "Alpha paragraph." in markdown
    assert len(chunks) == 2
    assert [chunk.text for chunk in chunks] == [
        "Alpha paragraph.",
        "Beta paragraph.",
    ]


@pytest.mark.anyio
async def test_convert_document_from_docx_file(tmp_path: Path) -> None:
    _assert_pandoc_available()

    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Alpha paragraph.")
    document.add_paragraph("Beta paragraph.")
    document.save(docx_path)

    markdown = await convert_document_to_markdown(docx_path)
    chunks = chunk_markdown(markdown)

    assert "Alpha paragraph." in markdown
    assert len(chunks) == 2
    assert [chunk.text for chunk in chunks] == [
        "Alpha paragraph.",
        "Beta paragraph.",
    ]

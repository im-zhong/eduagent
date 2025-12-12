from __future__ import annotations

from collections.abc import AsyncIterator
from pathlib import Path

import pytest
import pytest_asyncio
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from eduagent.documents.models import DocumentIngestionJob
from eduagent.documents.repository import DocumentRepository
from eduagent.documents.services import DocxIngestionService
from eduagent.user.models import Base

MAX_CHUNK_CHARS = 2048


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


@pytest_asyncio.fixture
async def session() -> AsyncIterator[AsyncSession]:
    engine = create_async_engine("sqlite+aiosqlite:///:memory:")
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    session_factory = async_sessionmaker(engine, expire_on_commit=False)
    async with session_factory() as session:
        yield session
    await engine.dispose()


@pytest.mark.asyncio
async def test_docx_ingestion_service_success(
    session: AsyncSession, tmp_path: Path
) -> None:
    file_path = tmp_path / "lesson.docx"
    _write_docx(
        file_path,
        [
            "Photosynthesis converts light energy into chemical energy.",
            "Chlorophyll absorbs light most efficiently in the blue portion.",
            "ATP and NADPH are produced in the light-dependent reactions.",
        ],
    )
    repo = DocumentRepository(session)
    service = DocxIngestionService(repo, chunk_size=40, chunk_overlap=0)

    job = await service.ingest_docx(
        source_filename="lesson.docx",
        file_path=str(file_path),
        subject="biology",
        grade_level="grade-8",
        metadata={"language": "en"},
    )
    assert job.status == "completed"
    assert job.total_chunks > 0

    chunks = await repo.list_chunks(job.id)
    assert len(chunks) == job.total_chunks
    assert any("paragraphs" in chunk.chunk_metadata for chunk in chunks)
    assert all(len(chunk.content) <= MAX_CHUNK_CHARS for chunk in chunks)


@pytest.mark.asyncio
async def test_docx_ingestion_service_handles_empty_doc(
    session: AsyncSession, tmp_path: Path
) -> None:
    empty_path = tmp_path / "empty.docx"
    _write_docx(empty_path, ["   "])
    repo = DocumentRepository(session)
    service = DocxIngestionService(repo)

    with pytest.raises(ValueError, match="No readable text"):
        await service.ingest_docx(
            source_filename="empty.docx",
            file_path=str(empty_path),
            subject=None,
            grade_level=None,
        )
    stmt = select(DocumentIngestionJob).where(
        DocumentIngestionJob.source_filename == "empty.docx"
    )
    result = await session.execute(stmt)
    failed_job = result.scalar_one()
    assert failed_job.status == "failed"
    assert failed_job.error_message is not None


@pytest.mark.asyncio
async def test_docx_ingestion_service_reads_table_content(
    session: AsyncSession, tmp_path: Path
) -> None:
    table_path = tmp_path / "table.docx"
    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Photosynthesis stages"
    table.cell(0, 1).text = "Light-dependent"
    table.cell(1, 0).text = "Calvin cycle"
    table.cell(1, 1).text = "Produces sugars"
    doc.save(str(table_path))

    repo = DocumentRepository(session)
    service = DocxIngestionService(repo, chunk_size=40, chunk_overlap=0)
    job = await service.ingest_docx(
        source_filename="table.docx",
        file_path=str(table_path),
        subject="science",
        grade_level="grade-7",
    )
    assert job.status == "completed"
    assert job.total_chunks > 0
    chunks = await repo.list_chunks(job.id)
    assert all(len(chunk.content) <= MAX_CHUNK_CHARS for chunk in chunks)


@pytest.mark.asyncio
async def test_docx_ingestion_service_ingests_real_doc(
    session: AsyncSession, tmp_path: Path
) -> None:
    source_path = Path("docs/classes.docx")
    # if not source_path.exists():  # pragma: no cover - defensive
    #     pytest.skip("docs/classes.docx not available")
    target = tmp_path / "classes.docx"
    target.write_bytes(source_path.read_bytes())

    repo = DocumentRepository(session)
    service = DocxIngestionService(repo)
    job = await service.ingest_docx(
        source_filename="classes.docx",
        file_path=str(target),
        subject="language",
        grade_level="reference",
        metadata={"source": "docs/classes.docx"},
    )
    assert job.status == "completed"
    assert job.total_chunks > 0
    chunks = await repo.list_chunks(job.id)
    assert chunks
    assert all(len(chunk.content) <= MAX_CHUNK_CHARS for chunk in chunks)


@pytest.mark.asyncio
async def test_document_repository_delete_all_data(
    session: AsyncSession, tmp_path: Path
) -> None:
    file_path = tmp_path / "cleanup.docx"
    _write_docx(file_path, ["示例课程内容", "第二段文本"])
    repo = DocumentRepository(session)
    service = DocxIngestionService(repo, chunk_size=40, chunk_overlap=0)
    job = await service.ingest_docx(
        source_filename="cleanup.docx",
        file_path=str(file_path),
        subject="general",
        grade_level="demo",
    )
    assert job.total_chunks > 0

    await repo.delete_all()

    remaining_jobs = await session.execute(select(DocumentIngestionJob))
    assert remaining_jobs.scalars().all() == []

"""Integration tests for document parse endpoint."""
from __future__ import annotations

from collections.abc import AsyncGenerator
from io import BytesIO

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from eduagent.api.api import api


@pytest.fixture(scope="function")
async def async_client(auth_token: str) -> AsyncGenerator[AsyncClient]:
    """Create async HTTP client for testing with authentication."""
    transport = ASGITransport(app=api, raise_app_exceptions=False)
    async with AsyncClient(
        transport=transport,
        base_url="http://test",
        headers={"Authorization": auth_token},
    ) as client:
        yield client


@pytest.mark.anyio
async def test_parse_document_creates_artifact_and_chunks(
    async_client: AsyncClient,
) -> None:
    """Test parsing a document returns artifact info and chunk count."""
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Alpha paragraph.")
    document.add_paragraph("Beta paragraph.")
    document.save(buffer)
    buffer.seek(0)

    upload_response = await async_client.post(
        "/api/v1/documents",
        files={"file": ("sample.docx", buffer.read())},
    )
    assert upload_response.status_code == 201
    doc_id = upload_response.json()["id"]

    parse_response = await async_client.post(
        "/api/v1/documents/parse",
        json={"doc_id": doc_id},
    )
    assert parse_response.status_code == 200
    payload = parse_response.json()
    assert payload["doc_id"] == doc_id
    assert payload["chunk_count"] == 2
    assert payload["artifact_id"] > 0
    assert payload["artifact_path"].endswith(".md")
